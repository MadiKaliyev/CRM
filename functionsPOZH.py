import os
from tkinter import messagebox
from docx import Document
from copy import deepcopy
from datetime import datetime

def set_cell_text(cell, text):
    cell.text = ''
    run = cell.paragraphs[0].add_run(text)
    cell.paragraphs[0].alignment = 1  # Центрирование текста

def replace_text(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                run.bold = True

def save_to_docs_POZH(company_name, people_data, chislo_chelovek):
    if not company_name:
        messagebox.showerror("Ошибка", "Пожалуйста, введите название компании.")
        return

    # Открытие документа ПОЖ.БЕЗ-НОМЕРА.docx
    file_path_1 = os.path.join(r"C:\Users\Madi\Desktop\ПРОГРАММА\Программа\НОМЕРА", "ПОЖ.БЕЗ-НОМЕРА.docx")
    try:
        doc1 = Document(file_path_1)
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при открытии файла ПОЖ.БЕЗ-НОМЕРА.docx: {e}")
        return

    table1 = doc1.tables[0]
    max_num_1 = 0
    max_num_4 = 0
    previous_people_end = 0

    current_year = datetime.now().year
    current_day = datetime.now().day
    current_month = datetime.now().month
    russian_months = {
        1: "января", 2: "февраля", 3: "марта", 4: "апреля", 5: "мая",
        6: "июня", 7: "июля", 8: "августа", 9: "сентября",
        10: "октября", 11: "ноября", 12: "декабря"
    }

    for row in table1.rows:
        if row.cells[0].text.isdigit():
            max_num_1 = max(max_num_1, int(row.cells[0].text))
        if row.cells[3].text.isdigit():
            max_num_4 = max(max_num_4, int(row.cells[3].text))
        previous_text = row.cells[4].text.split("-")
        if len(previous_text) == 2 and previous_text[1].isdigit():
            previous_people_end = int(previous_text[1])
        elif len(previous_text) == 1 and previous_text[0].isdigit():
            previous_people_end = int(previous_text[0])

    found_empty_row = False
    start_range = previous_people_end + 1  # Начальное значение диапазона
    end_range = start_range + len(people_data) - 1  # Конечное значение диапазона

    for row in table1.rows:
        if all(cell.text == '' for cell in row.cells):
            found_empty_row = True

            set_cell_text(row.cells[0], str(max_num_1 + 1))
            set_cell_text(row.cells[1], company_name)
            set_cell_text(row.cells[2], str(len(people_data)))
            set_cell_text(row.cells[3], str(max_num_4 + 1))

            if len(people_data) > 1:
                set_cell_text(row.cells[4], f"{start_range}-{end_range}")
            else:
                set_cell_text(row.cells[4], str(start_range))

            set_cell_text(row.cells[5], datetime.now().strftime("%d.%m.%Y"))
            break

    if not found_empty_row:
        new_row = table1.add_row()
        set_cell_text(new_row.cells[0], str(max_num_1 + 1))
        set_cell_text(new_row.cells[1], company_name)
        set_cell_text(new_row.cells[2], str(len(people_data)))
        set_cell_text(new_row.cells[3], str(max_num_4 + 1))

        if len(people_data) > 1:
            set_cell_text(new_row.cells[4], f"{start_range}-{end_range}")
        else:
            set_cell_text(new_row.cells[4], str(start_range))

        set_cell_text(new_row.cells[5], datetime.now().strftime("%d.%m.%Y"))

    doc1.save(file_path_1)

    # Открытие документа ПротоколПТМ.docx
    file_path_2 = os.path.join(r"C:\Users\Madi\Desktop\ПРОГРАММА\Программа\файлы", "ПротоколПТМ.docx")
    try:
        doc2 = Document(file_path_2)
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при открытии файла ПротоколПТМ.docx: {e}")
        return

    num_protocol = max_num_4 + 1
    current_date = datetime.now().strftime("%d.%m.%Y")

    for paragraph in doc2.paragraphs:
        replace_text(paragraph, 'НОМЕР', str(num_protocol))
        replace_text(paragraph, 'ДАТА', f'{current_day} {russian_months[current_month]} {current_year}')

    table2 = doc2.tables[0]
    max_num_2 = 0
    for row in table2.rows:
        if row.cells[0].text.isdigit():
            max_num_2 = max(max_num_2, int(row.cells[0].text))

    for fio, position in people_data:
        found_empty_row = False
        for row in table2.rows:
            if all(cell.text == '' for cell in row.cells):
                found_empty_row = True
                set_cell_text(row.cells[0], str(max_num_2 + 1))
                set_cell_text(row.cells[1], fio)
                set_cell_text(row.cells[2], position)
                set_cell_text(row.cells[3], company_name)
                set_cell_text(row.cells[4], "Очередная")
                set_cell_text(row.cells[5], "Сдал")
                max_num_2 += 1
                break

        if not found_empty_row:
            new_row = table2.add_row()
            set_cell_text(new_row.cells[0], str(max_num_2 + 1))
            set_cell_text(new_row.cells[1], fio)
            set_cell_text(new_row.cells[2], position)
            set_cell_text(new_row.cells[3], company_name)
            set_cell_text(new_row.cells[4], "Очередная")
            set_cell_text(new_row.cells[5], "Сдал")
            max_num_2 += 1

    sanitized_company_name = company_name.replace('"', '').replace(":", "").replace("?", "").replace("*", "")
    predpriyatiya_path = os.path.join(r"C:\Users\Madi\Desktop", "ПРЕДПРИЯТИЯ")

    # Путь к основной папке с именем sanitized_company_name
    company_directory = os.path.join(predpriyatiya_path, sanitized_company_name)
    if not os.path.exists(company_directory):
        os.makedirs(company_directory)

    # Создаем уникальную поддиректорию с датой и названием компании
    base_name = f"{sanitized_company_name} {current_day} {russian_months[current_month]} {current_year}"
    main_directory = os.path.join(company_directory, base_name)

    if not os.path.exists(main_directory):
        os.makedirs(main_directory)

    # Внутри main_directory создаем поддиректорию "Пож.Без"
    sub_directory = os.path.join(main_directory, "Пож.Без")
    counter = 1
    while os.path.exists(sub_directory):
        sub_directory = os.path.join(main_directory, f"Пож.Без({counter})")
        counter += 1
    os.makedirs(sub_directory)

    # Сохраняем документ doc2 в поддиректории "Пож.Без"
    try:
        doc2.save(os.path.join(sub_directory, "Протокол.docx"))
        messagebox.showinfo("Успех", f"Данные успешно записаны в файлы в папке {sub_directory}.")
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при сохранении файла: {e}")

    # Открытие документа ШАБЛОНПТМ.docx
    file_path_template = os.path.join(r"C:\Users\Madi\Desktop\ПРОГРАММА\Программа\файлы", "ШАБЛОНПТМ.docx")
    try:
        doc_template = Document(file_path_template)
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при открытии файла ШАБЛОНПТМ.docx: {e}")
        return

    first_table = doc_template.tables[0]

    if len(people_data) > 1:
        for i in range(1, len(people_data)):
            new_table = deepcopy(first_table)
            doc_template.element.body.append(new_table._element)

    for paragraph in doc_template.paragraphs:
        if not paragraph.text.strip():
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

    start_range = previous_people_end + 1
    for i, table in enumerate(doc_template.tables):
        if i >= len(people_data):
            break

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fio_parts = people_data[i][0].split(' ', 1)
                    first_name = fio_parts[0]
                    last_name = fio_parts[1] if len(fio_parts) > 1 else ""

                    replace_text(paragraph, 'ФИО', first_name)
                    replace_text(paragraph, 'ЧАСТЬДВА', last_name)
                    replace_text(paragraph, 'НОМЕРУДВ', str(start_range))
                    replace_text(paragraph, 'КОМПАНИЯ', company_name)
                    replace_text(paragraph, 'ДОЛЖНОСТЬ', people_data[i][1])

                    if i < int(chislo_chelovek):
                        replace_text(paragraph, 'датаказ', f"{current_day} {russian_months[current_month]} {current_year + 1}")
                        replace_text(paragraph, 'ДАДАКОНЕЦ', f"{current_day} {russian_months[current_month]} {current_year + 1}")
                    else:
                        replace_text(paragraph, 'датаказ', f"{current_day} {russian_months[current_month]} {current_year + 3}")
                        replace_text(paragraph, 'ДАДАКОНЕЦ', f"{current_day} {russian_months[current_month]} {current_year + 3}")

                    replace_text(paragraph, 'МЕСЯЦ', russian_months[current_month])
                    replace_text(paragraph, 'ГОД', str(current_year))
                    replace_text(paragraph, 'ДАТА', f"{current_day} {russian_months[current_month]} {current_year}")
                    replace_text(paragraph, 'НОМЕР', str(num_protocol))
            start_range += 1

    try:
        doc_template.save(os.path.join(sub_directory, "Удостоверение.docx"))
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при сохранении файла Удостоверение: {e}")
