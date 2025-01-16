import os
from tkinter import messagebox
from docx import Document
from copy import deepcopy
from datetime import datetime

def set_cell_text(cell, text):
    cell.text = ''
    run = cell.paragraphs[0].add_run(text)
    cell.paragraphs[0].alignment = 1  # Центрирование текста

def replace_text(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                run.bold = True

def save_to_docs_biot(company_name, people_data, chislo_chelovek):
    if not company_name:
        messagebox.showerror("Ошибка", "Пожалуйста, введите название компании.")
        return

    # Открытие документа БиОТ.РАБ-НОМЕРА.docx
    file_path_1 = r"C:\Users\\Madi\Desktop\ПРОГРАММА\Программа\НОМЕРА\БиОТ.РАБ-НОМЕРА.docx"
    try:
        doc1 = Document(file_path_1)
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при открытии файла БиОТ.РАБ-НОМЕРА.docx: {e}")
        return

    table1 = doc1.tables[0]
    max_num_1 = 0
    max_num_4 = 0
    previous_people_end = 0

    current_year = datetime.now().year
    current_day = datetime.now().day
    current_month = datetime.now().month
    russian_months = {
        1: "января", 2: "февраля", 3: "марта", 4: "апреля", 5: "мая",
        6: "июня", 7: "июля", 8: "августа", 9: "сентября",
        10: "октября", 11: "ноября", 12: "декабря"
    }

    for row in table1.rows:
        if row.cells[0].text.isdigit():
            max_num_1 = max(max_num_1, int(row.cells[0].text))
        if row.cells[3].text.isdigit():
            max_num_4 = max(max_num_4, int(row.cells[3].text))
        previous_text = row.cells[4].text.split("-")
        if len(previous_text) == 2 and previous_text[1].isdigit():
            previous_people_end = int(previous_text[1])
        elif len(previous_text) == 1 and previous_text[0].isdigit():
            previous_people_end = int(previous_text[0])

    start_range = previous_people_end + 1  # Начальное значение диапазона
    end_range = start_range + len(people_data) - 1  # Конечное значение диапазона

    found_empty_row = False
    for row in table1.rows:
        if all(cell.text == '' for cell in row.cells):
            found_empty_row = True

            set_cell_text(row.cells[0], str(max_num_1 + 1))
            set_cell_text(row.cells[1], company_name)
            set_cell_text(row.cells[2], str(len(people_data)))
            set_cell_text(row.cells[3], str(max_num_4 + 1))

            if len(people_data) > 1:
                set_cell_text(row.cells[4], f"{start_range}-{end_range}")
            else:
                set_cell_text(row.cells[4], str(start_range))

            set_cell_text(row.cells[5], datetime.now().strftime("%d.%m.%Y"))
            break

    if not found_empty_row:
        new_row = table1.add_row()
        set_cell_text(new_row.cells[0], str(max_num_1 + 1))
        set_cell_text(new_row.cells[1], company_name)
        set_cell_text(new_row.cells[2], str(len(people_data)))
        set_cell_text(new_row.cells[3], str(max_num_4 + 1))

        if len(people_data) > 1:
            set_cell_text(new_row.cells[4], f"{start_range}-{end_range}")
        else:
            set_cell_text(new_row.cells[4], str(start_range))

        set_cell_text(new_row.cells[5], datetime.now().strftime("%d.%m.%Y"))

    doc1.save(file_path_1)

    # Открытие документа ПротоколБИОТраб.docx
    file_path_2 = r"C:\Users\\Madi\Desktop\ПРОГРАММА\Программа\файлы\ПротоколБИОТраб.docx"
    try:
        doc2 = Document(file_path_2)
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при открытии файла ПротоколБИОТраб.docx: {e}")
        return

    num_protocol = max_num_4 + 1

    for paragraph in doc2.paragraphs:
        replace_text(paragraph, 'НОМЕР', str(num_protocol))
        replace_text(paragraph, 'ДАТА', f'{current_day} {russian_months[current_month]} {current_year}')

    table2 = doc2.tables[0]
    max_num_2 = 0
    for row in table2.rows:
        if row.cells[0].text.isdigit():
            max_num_2 = max(max_num_2, int(row.cells[0].text))

    for fio, position in people_data:
        found_empty_row = False
        for row in table2.rows:
            if all(cell.text == '' for cell in row.cells):
                found_empty_row = True
                set_cell_text(row.cells[0], str(max_num_2 + 1))
                set_cell_text(row.cells[1], fio)
                set_cell_text(row.cells[2], company_name)
                set_cell_text(row.cells[3], position)
                set_cell_text(row.cells[4], "Сдал")
                set_cell_text(row.cells[5], str(start_range))
                max_num_2 += 1
                break

        if not found_empty_row:
            new_row = table2.add_row()
            set_cell_text(new_row.cells[0], str(max_num_2 + 1))
            set_cell_text(new_row.cells[1], fio)
            set_cell_text(new_row.cells[2], company_name)
            set_cell_text(new_row.cells[3], position)
            set_cell_text(new_row.cells[4], "Сдал")
            set_cell_text(new_row.cells[5], f'№ {str(start_range)}')
            start_range += 1
            max_num_2 += 1

    sanitized_company_name = company_name.replace('"', '').replace(":", "").replace("?", "").replace("*", "")
    predpriyatiya_path = r"C:\Users\\Madi\Desktop\ПРЕДПРИЯТИЯ"

    # Путь к основной папке с именем sanitized_company_name
    company_directory = os.path.join(predpriyatiya_path, sanitized_company_name)
    if not os.path.exists(company_directory):
        os.makedirs(company_directory)

    # Создаем поддиректорию с датой и названием компании
    base_name = f"{sanitized_company_name} {current_day} {russian_months[current_month]} {current_year}"
    main_directory = os.path.join(company_directory, base_name)

    # Проверяем, существует ли main_directory
    if not os.path.exists(main_directory):
        os.makedirs(main_directory)
    else:
        # Если main_directory уже существует, используем её
        pass

    # Внутри main_directory создаем поддиректорию "БиОТ.Раб"
    sub_directory = os.path.join(main_directory, "БиОТ.Раб")
    counter = 1
    while os.path.exists(sub_directory):
        sub_directory = os.path.join(main_directory, f"БиОТ.Раб({counter})")
        counter += 1
    os.makedirs(sub_directory)

    # Сохраняем документ doc2 в поддиректории "БиОТ.Раб"
    doc2.save(os.path.join(sub_directory, "Протокол.docx"))

    # Работа с шаблоном ШАБЛОНБИОТ.docx
    file_path_template = r"C:\Users\Madi\Desktop\ПРОГРАММА\Программа\файлы\ШАБЛОНБИОТ.docx"
    try:
        doc_template = Document(file_path_template)
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при открытии файла ШАБЛОНБИОТ.docx: {e}")
        return

    current_year_over = current_year + 1
    current_year_over3 = current_year + 3

    kazakh_months = {
        1: "қаңтарда", 2: "ақпанда", 3: "наурызда", 4: "сәуірде", 5: "мамырда",
        6: "маусымда", 7: "шілдеде", 8: "тамызда", 9: "қыркүйекте",
        10: "қазанда", 11: "қарашада", 12: "желтоқсанда"
    }

    kazakh_months2 = {
        1: "қаңтар", 2: "ақпан", 3: "наурыз", 4: "сәуір", 5: "мамыр",
        6: "маусым", 7: "шілде", 8: "тамыз", 9: "қыркүйек",
        10: "қазан", 11: "қараша", 12: "желтоқсан"
    }

    current_month_kazakh = kazakh_months[current_month]
    current_month_kazakh2 = kazakh_months2[current_month]

    first_table = doc_template.tables[0]

    if len(people_data) > 1:
        for i in range(1, len(people_data)):
            new_table = deepcopy(first_table)
            doc_template.element.body.append(new_table._element)

    for paragraph in doc_template.paragraphs:
        if not paragraph.text.strip():
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

    start_range = previous_people_end + 1  # Обновляем начальное значение диапазона для шаблона

    for i, table in enumerate(doc_template.tables):
        if i >= len(people_data):
            break

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fio_parts = people_data[i][0].split(' ', 1)
                    first_name = fio_parts[0]
                    last_name = fio_parts[1] if len(fio_parts) > 1 else ""

                    replace_text(paragraph, 'ФИО', people_data[i][0])
                    replace_text(paragraph, 'НОМЕРУДВ', str(start_range))
                    replace_text(paragraph, 'КОМПАНИЯ', company_name)
                    replace_text(paragraph, 'ДОЛЖНОСТЬ', people_data[i][1])

                    if i < int(chislo_chelovek):
                        replace_text(paragraph, 'датаказ',
                                     f"{current_day} {current_month_kazakh2} {current_year_over}")
                        replace_text(paragraph, 'ДАДАКОНЕЦ',
                                     f"{current_day} {russian_months[current_month]} {current_year_over}")
                        replace_text(paragraph, 'ЧАС', str(10))
                    else:
                        replace_text(paragraph, 'датаказ',
                                     f"{current_day} {current_month_kazakh2} {current_year_over3}")
                        replace_text(paragraph, 'ДАДАКОНЕЦ',
                                     f"{current_day} {russian_months[current_month]} {current_year_over3}")
                        replace_text(paragraph, 'ЧАС', str(40))

                    replace_text(paragraph, 'МЕСЯЦ', russian_months[current_month])
                    replace_text(paragraph, 'ГОД', str(current_year))
                    replace_text(paragraph, 'ДАТА', f"{current_day} {russian_months[current_month]} {current_year}")
                    replace_text(paragraph, 'НОМЕР', str(num_protocol))

        start_range += 1

    doc_template.save(rf"{sub_directory}/Удостоверение.docx")
    messagebox.showinfo("Успех", f"Данные успешно записаны в файлы в папке {sub_directory}.")
