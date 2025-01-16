import os
from tkinter import messagebox
from docx import Document
from copy import deepcopy
from datetime import datetime
from win32com import client
import re

# Функция для конвертации .docx в .pdf
def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        word = client.Dispatch("Word.Application")
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 соответствует формату PDF
        doc.Close()
        word.Quit()
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при конвертации в PDF: {e}")

# Функция для установки текста в ячейку таблицы
def set_cell_text(cell, text):
    cell.text = ''
    run = cell.paragraphs[0].add_run(text)
    cell.paragraphs[0].alignment = 1  # Центрирование текста

# Функция для замены текста в параграфах
def replace_text(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                run.bold = True

# Функция для замены текста в таблице
def replace_text_in_table(table, old_text, new_text):
    """Замена текста в таблице."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text(paragraph, old_text, new_text)

# Функция для очистки имени файла от недопустимых символов
def sanitize_filename(filename):
    return re.sub(r'[\/:*?"<>|\\\n]', '_', filename)

# Основная функция для сохранения данных
def save_to_docs_biot2(company_name, people_data, chislo_chelovek):
    if not company_name:
        messagebox.showerror("Ошибка", "Пожалуйста, введите название компании.")
        return

    # Открытие документа БиОТ.РУКОВОДИТЕЛИ.docx
    file_path_1 = r"C:\Users\Madi\Desktop\ПРОГРАММА\Программа\НОМЕРА\БиОТ.РУКОВОДИТЕЛИ.docx"
    try:
        doc1 = Document(file_path_1)
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при открытии файла БиОТ.РУКОВОДИТЕЛИ.docx: {e}")
        return

    table1 = doc1.tables[0]
    max_num_1 = 0
    max_num_4 = 0
    previous_people_end = 0

    current_year = datetime.now().year
    current_day = datetime.now().day
    current_month = datetime.now().month

    russian_months = {
        1: "января", 2: "февраля", 3: "марта", 4: "апреля", 5: "мая",
        6: "июня", 7: "июля", 8: "августа", 9: "сентября",
        10: "октября", 11: "ноября", 12: "декабря"
    }

    kazakh_months = {
        1: "қаңтар", 2: "ақпан", 3: "наурыз", 4: "сәуір", 5: "мамыр",
        6: "маусым", 7: "шілде", 8: "тамыз", 9: "қыркүйек",
        10: "қазан", 11: "қараша", 12: "желтоқсан"
    }

    # Подсчет номеров для протокола
    for row in table1.rows:
        if row.cells[0].text.isdigit():
            max_num_1 = max(max_num_1, int(row.cells[0].text))
        if row.cells[3].text.isdigit():
            max_num_4 = max(max_num_4, int(row.cells[3].text))
        previous_text = row.cells[4].text.split("-")
        if len(previous_text) == 2 and previous_text[1].isdigit():
            previous_people_end = int(previous_text[1])
        elif len(previous_text) == 1 and previous_text[0].isdigit():
            previous_people_end = int(previous_text[0])

    start_range = previous_people_end + 1
    end_range = start_range + len(people_data) - 1

    found_empty_row = False
    for row in table1.rows:
        if all(cell.text == '' for cell in row.cells):
            found_empty_row = True

            set_cell_text(row.cells[0], str(max_num_1 + 1))
            set_cell_text(row.cells[1], company_name)
            set_cell_text(row.cells[2], str(len(people_data)))
            set_cell_text(row.cells[3], str(max_num_4 + 1))

            if len(people_data) > 1:
                set_cell_text(row.cells[4], f"{start_range}-{end_range}")
            else:
                set_cell_text(row.cells[4], str(start_range))

            set_cell_text(row.cells[5], datetime.now().strftime("%d.%m.%Y"))
            break

    if not found_empty_row:
        new_row = table1.add_row()
        set_cell_text(new_row.cells[0], str(max_num_1 + 1))
        set_cell_text(new_row.cells[1], company_name)
        set_cell_text(new_row.cells[2], str(len(people_data)))
        set_cell_text(new_row.cells[3], str(max_num_4 + 1))

        if len(people_data) > 1:
            set_cell_text(new_row.cells[4], f"{start_range}-{end_range}")
        else:
            set_cell_text(new_row.cells[4], str(start_range))

        set_cell_text(new_row.cells[5], datetime.now().strftime("%d.%m.%Y"))

    doc1.save(file_path_1)

    # Открытие документа ПротоколБИОТраб.docx
    file_path_2 = r"C:\Users\Madi\Desktop\ПРОГРАММА\Программа\файлы\БИОТ.docx"
    try:
        doc2 = Document(file_path_2)
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при открытии файла БИОТ.docx: {e}")
        return

    num_protocol = max_num_4 + 1

    for paragraph in doc2.paragraphs:
        replace_text(paragraph, 'НОМЕР', str(num_protocol))
        replace_text(paragraph, 'ДАТА', f'{current_day} {russian_months[current_month]} {current_year}')

    table2 = doc2.tables[0]
    max_num_2 = 0
    for row in table2.rows:
        if row.cells[0].text.isdigit():
            max_num_2 = max(max_num_2, int(row.cells[0].text))

    for fio, position in people_data:
        sanitized_fio = sanitize_filename(fio)
        found_empty_row = False
        for row in table2.rows:
            if all(cell.text == '' for cell in row.cells):
                found_empty_row = True
                set_cell_text(row.cells[0], str(max_num_2 + 1))
                set_cell_text(row.cells[1], fio)
                set_cell_text(row.cells[2], company_name)
                set_cell_text(row.cells[3], position)
                set_cell_text(row.cells[4], "Сдал")
                set_cell_text(row.cells[5], str(start_range))
                max_num_2 += 1
                break

        if not found_empty_row:
            new_row = table2.add_row()
            set_cell_text(new_row.cells[0], str(max_num_2 + 1))
            set_cell_text(new_row.cells[1], fio)
            set_cell_text(new_row.cells[2], company_name)
            set_cell_text(new_row.cells[3], position)
            set_cell_text(new_row.cells[4], "Сдал")
            set_cell_text(new_row.cells[5], f'№ {str(start_range)}')
            start_range += 1
            max_num_2 += 1

    sanitized_company_name = sanitize_filename(company_name)
    predpriyatiya_path = r"C:\Users\Madi\Desktop\ПРЕДПРИЯТИЯ"

    # Путь к основной папке с именем sanitized_company_name
    company_directory = os.path.join(predpriyatiya_path, sanitized_company_name)
    if not os.path.exists(company_directory):
        os.makedirs(company_directory)

    # Создаем поддиректорию с датой и названием компании
    base_name = f"{sanitized_company_name} {current_day} {russian_months[current_month]} {current_year}"
    main_directory = os.path.join(company_directory, base_name)

    # Проверяем, существует ли main_directory
    if not os.path.exists(main_directory):
        os.makedirs(main_directory)

    # Внутри main_directory создаем поддиректорию "БиОТ"
    sub_directory = os.path.join(main_directory, "БиОТ")
    counter = 1
    while os.path.exists(sub_directory):
        sub_directory = os.path.join(main_directory, f"БиОТ({counter})")
        counter += 1
    os.makedirs(sub_directory)

    # Сохраняем документ doc2 в поддиректории "БиОТ"
    doc2.save(os.path.join(sub_directory, "Протокол.docx"))

    # Открытие шаблона СЕРТИФИКАТ
    template_path = r"C:\Users\Madi\Desktop\ПРОГРАММА\Программа\файлы\ШАБЛОНСЕРТИФИКАТ.docx"
    try:
        doc_template = Document(template_path)
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при открытии файла ШАБЛОНСЕРТИФИКАТ.docx: {e}")
        return

    start = previous_people_end + 1

    # Создаем сертификат для каждого человека
    for idx, (fio, position) in enumerate(people_data):
        current_number = start + idx
        sanitized_fio = sanitize_filename(fio)

        doc_copy = deepcopy(doc_template)

        # Замена текста в параграфах
        for paragraph in doc_copy.paragraphs:
            replace_text(paragraph, "ФИО", fio)
            replace_text(paragraph, "ДАТА", f"{current_day} {kazakh_months[current_month]} {current_year}")
            replace_text(paragraph, "ДАТ", f"{current_day} {russian_months[current_month]} {current_year}")
            replace_text(paragraph, "000", str(current_number))
            replace_text(paragraph, "001", str(current_number))

        # Замена текста в таблицах
        for table in doc_copy.tables:
            replace_text_in_table(table, "ФИО", fio)
            replace_text_in_table(table, "ДАТА", f"{current_day} {kazakh_months[current_month]} {current_year}")
            replace_text_in_table(table, "ДАТ", f"{current_day} {russian_months[current_month]} {current_year}")
            replace_text_in_table(table, "000", str(current_number))
            replace_text_in_table(table, "001", str(current_number))

        # Убедимся, что для каждого человека создается отдельный сертификат с уникальным именем
        docx_path = os.path.join(sub_directory, f"СЕРТИФИКАТ_{sanitized_fio}_{current_number}.docx")
        doc_copy.save(docx_path)

        # Путь для сохранения PDF
        pdf_path = os.path.join(sub_directory, f"СЕРТИФИКАТ_{sanitized_fio}_{current_number}.pdf")

        # Конвертация docx в pdf
        convert_docx_to_pdf(docx_path, pdf_path)

    messagebox.showinfo("Успех", f"Данные успешно записаны в файлы в папке {sub_directory}.")
